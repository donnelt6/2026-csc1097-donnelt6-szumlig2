"""tasks.py: Compatibility facade plus worker task orchestration."""

import hashlib
import json
import re
import ssl
import uuid
from collections import defaultdict
from datetime import datetime, timedelta, timezone
from typing import Iterable, List, Optional
from urllib.parse import parse_qs, urlencode, urlparse, urlunparse
from zoneinfo import ZoneInfo

import dateparser
import redis
import spacy
from openai import OpenAI
from supabase import Client

from . import common as _common
from . import content as _content
from . import response_utils as _response_utils
from . import storage as _storage
from . import web as _web
from . import youtube as _youtube
from .app import celery_app, logger, settings


# Celery ingestion tasks.
# These task entrypoints stay in `worker.tasks` so existing Celery commands
# and task names keep working while helper logic lives in split modules.
# Ingests an uploaded file source by downloading, extracting, chunking, and storing its content.
@celery_app.task(bind=True, name="ingest_source", max_retries=3, default_retry_delay=15)
def ingest_source(self, source_id: str, hub_id: str, storage_path: str) -> dict:
    """
    Ingestion flow:
    - Download from Supabase Storage
    - Extract text (PDF/DOCX/TXT/MD)
    - Chunk + embed
    - Persist chunks to pgvector
    - Update source status
    """
    logger.info("worker.ingest.start source_id=%s", source_id)
    client = _get_supabase_client()
    _update_source(client, source_id, status="processing", clear_failure_reason=True)

    try:
        raw = _download_from_storage(storage_path)
    except Exception as exc:
        logger.warning("worker.ingest.download_retry storage_path=%s error=%s", storage_path, exc)
        raise self.retry(exc=exc)

    try:
        text = _extract_text(raw, storage_path)
        text = _normalize_text(text)
        if not text:
            raise ValueError("No text extracted from source")

        chunk_count = _ingest_text_for_source(client, source_id, hub_id, text, extra_metadata=None)
        logger.info("worker.ingest.complete source_id=%s", source_id)
        return {"source_id": source_id, "hub_id": hub_id, "chunks": chunk_count}
    except Exception as exc:
        logger.exception("worker.ingest.failed source_id=%s", source_id)
        _update_source(client, source_id, status="failed", failure_reason=str(exc)[:500])
        raise


# Ingests a web source by fetching the page, extracting text, and storing the resulting chunks.
@celery_app.task(bind=True, name="ingest_web_source", max_retries=3, default_retry_delay=15)
def ingest_web_source(self, source_id: str, hub_id: str, url: str, storage_path: str) -> dict:
    """
    Web ingestion flow:
    - Validate URL (public only)
    - Respect robots.txt if enabled
    - Fetch + extract readable text
    - Store pseudo document in Supabase Storage
    - Chunk + embed + persist to pgvector
    """
    logger.info("worker.web_ingest.start source_id=%s url=%s", source_id, url)
    client = _get_supabase_client()
    _update_source(client, source_id, status="processing", clear_failure_reason=True)

    try:
        safe_url = _validate_public_url(url)
        if settings.web_respect_robots and not _allowed_by_robots(safe_url, settings.web_user_agent):
            raise ValueError("Blocked by robots.txt")
        raw, content_type, final_url = _fetch_url_content(safe_url)
        text, title = _extract_web_text(raw, content_type)
        text = _normalize_text(text)
        if not text:
            raise ValueError("No text extracted from web page")
        crawl_at = datetime.now(timezone.utc).isoformat()
        pseudo_doc = _build_pseudo_doc(title, final_url or safe_url, crawl_at, content_type, text)
        _upload_pseudo_doc(client, storage_path, pseudo_doc)
        extra_metadata = {
            "source_type": "web",
            "url": safe_url,
            "final_url": final_url,
            "title": title,
            "crawl_at": crawl_at,
            "content_type": content_type,
            "byte_size": len(raw),
            "word_count": len(text.split()),
        }
        chunk_count = _ingest_text_for_source(client, source_id, hub_id, text, extra_metadata=extra_metadata)
        try:
            if title:
                client.table("sources").update({"original_name": title[:500]}).eq("id", source_id).execute()
        except Exception:
            logger.warning("worker.web_ingest.title_update_failed source_id=%s", source_id, exc_info=True)
        logger.info("worker.web_ingest.complete source_id=%s", source_id)
        return {"source_id": source_id, "hub_id": hub_id, "chunks": chunk_count}
    except Exception as exc:
        logger.exception("worker.web_ingest.failed source_id=%s", source_id)
        _update_source(client, source_id, status="failed", failure_reason=str(exc)[:500])
        raise


# Ingests a YouTube source by downloading captions and storing the transcript content.
@celery_app.task(bind=True, name="ingest_youtube_source", max_retries=3, default_retry_delay=15)
def ingest_youtube_source(
    self,
    source_id: str,
    hub_id: str,
    url: str,
    storage_path: str,
    language: Optional[str] = None,
    allow_auto_captions: Optional[bool] = None,
    video_id: Optional[str] = None,
) -> dict:
    """
    YouTube ingestion flow:
    - Fetch video info + captions via yt-dlp
    - Store pseudo document in Supabase Storage
    - Chunk + embed + persist to pgvector
    """
    logger.info("worker.youtube_ingest.start source_id=%s url=%s", source_id, url)
    client = _get_supabase_client()
    _update_source(client, source_id, status="processing", clear_failure_reason=True)

    try:
        transcript, info, captions_meta = _fetch_youtube_transcript(
            url,
            language=language,
            allow_auto_captions=allow_auto_captions,
        )
        if video_id:
            info["video_id"] = video_id
        text = _normalize_text(transcript)
        if not text:
            raise ValueError("No transcript text extracted from YouTube captions")
        fetched_at = datetime.now(timezone.utc).isoformat()
        pseudo_doc = _build_youtube_pseudo_doc(info, url, fetched_at, captions_meta, text)
        _upload_pseudo_doc(client, storage_path, pseudo_doc)
        extra_metadata = {
            "source_type": "youtube",
            "url": url,
            "video_id": info.get("video_id"),
            "title": info.get("title"),
            "channel": info.get("channel"),
            "channel_id": info.get("channel_id"),
            "published_at": info.get("published_at"),
            "duration_seconds": info.get("duration_seconds"),
            "language": captions_meta.get("language"),
            "captions_source": captions_meta.get("captions_source"),
            "transcript_fetched_at": fetched_at,
            "word_count": len(text.split()),
        }
        chunk_count = _ingest_text_for_source(client, source_id, hub_id, text, extra_metadata=extra_metadata)
        try:
            video_title = info.get("title")
            if video_title:
                client.table("sources").update({"original_name": video_title[:500]}).eq("id", source_id).execute()
        except Exception:
            logger.warning("worker.youtube_ingest.title_update_failed source_id=%s", source_id, exc_info=True)
        logger.info("worker.youtube_ingest.complete source_id=%s", source_id)
        return {"source_id": source_id, "hub_id": hub_id, "chunks": chunk_count}
    except Exception as exc:
        logger.exception("worker.youtube_ingest.failed source_id=%s", source_id)
        _update_source(client, source_id, status="failed", failure_reason=str(exc)[:500])
        raise


# Scans eligible hubs and generates any new source suggestions that should be queued.
@celery_app.task(name="scan_source_suggestions")
def scan_source_suggestions() -> dict:
    client = _get_supabase_client()
    now = datetime.now(timezone.utc)
    eligible_hubs = _list_eligible_source_suggestion_hubs(client, now=now)
    processed = 0
    generated = 0

    for hub in eligible_hubs:
        hub_id = str(hub.get("id") or "")
        if not hub_id:
            continue
        # Use a short-lived Redis lock so multiple workers do not scan the same hub together.
        lock = _acquire_source_suggestion_lock(hub_id)
        if lock is None:
            logger.info("worker.source_suggestions.lock_held hub_id=%s", hub_id)
            continue
        try:
            result = _generate_source_suggestions_for_hub(client, hub_id, now=now)
            processed += 1
            generated += int(result.get("inserted", 0) or 0)
        except Exception:
            logger.exception("worker.source_suggestions.failed hub_id=%s", hub_id)
            _mark_source_suggestion_scan(client, hub_id, now=now, generated=False)
        finally:
            _release_source_suggestion_lock(lock)

    return {"eligible_hubs": len(eligible_hubs), "processed_hubs": processed, "generated": generated}


# Shared storage and source-state helpers.
# Creates a Supabase service client for worker-side database and storage operations.
def _get_supabase_client() -> Client:

    if not settings.supabase_url or not settings.supabase_service_role_key:
        raise RuntimeError("Supabase credentials missing in worker environment")
    return create_client(settings.supabase_url, settings.supabase_service_role_key)


# Downloads the raw file bytes for a source from Supabase Storage.
def _download_from_storage(storage_path: str) -> bytes:
    if not settings.supabase_url or not settings.supabase_service_role_key:
        raise RuntimeError("Supabase credentials missing for storage download")
    # Use service key to fetch the object directly from Supabase Storage.
    safe_path = quote(storage_path, safe="/")
    storage_url = f"{settings.supabase_url}/storage/v1/object/{settings.storage_bucket}/{safe_path}"
    headers = {
        "Authorization": f"Bearer {settings.supabase_service_role_key}",
        "apikey": settings.supabase_service_role_key,
    }
    with httpx.Client(timeout=60) as client:
        resp = client.get(storage_url, headers=headers)
        resp.raise_for_status()
    return resp.content


# Web fetching and validation helpers.
# Validates a URL and rejects unsupported schemes or private-network targets.
def _validate_public_url(url: str) -> str:

    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        raise ValueError("URL scheme must be http or https")
    if not parsed.netloc:
        raise ValueError("URL must include a host")
    hostname = parsed.hostname or ""
    if not hostname:
        raise ValueError("URL must include a host")
    _ensure_public_host(hostname)
    return parsed.geturl()


# Resolves a hostname and rejects loopback, private, or otherwise unsafe IP addresses.
def _ensure_public_host(hostname: str) -> None:
    ip_list: List[ipaddress.IPv4Address | ipaddress.IPv6Address] = []
    try:
        ip_list.append(ipaddress.ip_address(hostname))
    except ValueError:
        try:
            infos = socket.getaddrinfo(hostname, None)
        except OSError as exc:
            raise ValueError("Unable to resolve host") from exc
        for info in infos:
            addr = info[4][0]
            try:
                ip_list.append(ipaddress.ip_address(addr))
            except ValueError:
                continue
    for addr in ip_list:
        if addr.is_private or addr.is_loopback or addr.is_link_local or addr.is_reserved or addr.is_multicast or addr.is_unspecified:
            raise ValueError("URL resolves to a private or non-public address")


# Checks whether the worker is allowed to crawl the target URL under robots.txt rules.
def _allowed_by_robots(url: str, user_agent: str) -> bool:
    parsed = urlparse(url)
    robots_url = f"{parsed.scheme}://{parsed.netloc}/robots.txt"
    try:
        with httpx.Client(timeout=settings.web_timeout_seconds) as client:
            resp = client.get(robots_url, headers={"User-Agent": user_agent}, follow_redirects=True)
            if resp.status_code >= 400:
                return True
            parser = RobotFileParser()
            parser.parse(resp.text.splitlines())
            return parser.can_fetch(user_agent, url)
    except Exception:
        return True


# Fetches web content while enforcing the worker's size and timeout limits.
def _fetch_url_content(url: str) -> tuple[bytes, str, str]:
    headers = {"User-Agent": settings.web_user_agent}
    max_bytes = max(1, settings.web_max_bytes)
    current_url = url
    max_redirects = 5
    with httpx.Client(timeout=settings.web_timeout_seconds, follow_redirects=False) as client:
        for _ in range(max_redirects + 1):
            with client.stream("GET", current_url, headers=headers) as resp:
                if 300 <= resp.status_code < 400:
                    location = resp.headers.get("location")
                    if not location:
                        raise ValueError("Redirect without location header")
                    next_url = urljoin(current_url, location)
                    parsed = urlparse(next_url)
                    if not parsed.scheme or not parsed.netloc:
                        raise ValueError("Invalid redirect URL")
                    _ensure_public_host(parsed.hostname or "")
                    current_url = next_url
                    continue
                resp.raise_for_status()
                content_type = resp.headers.get("content-type", "")
                total = 0
                chunks: list[bytes] = []
                for chunk in resp.iter_bytes():
                    total += len(chunk)
                    if total > max_bytes:
                        raise ValueError("Web content exceeds size limit")
                    chunks.append(chunk)
                return b"".join(chunks), content_type, current_url
    raise ValueError("Too many redirects")


# Extracts readable text and an optional title from fetched web content.
def _extract_web_text(raw: bytes, content_type: str) -> tuple[str, Optional[str]]:
    encoding = "utf-8"
    match = re.search(r"charset=([\w-]+)", content_type, re.IGNORECASE)
    if match:
        encoding = match.group(1)
    html = raw.decode(encoding, errors="ignore")
    lowered = content_type.lower()
    if "text/html" not in lowered and "application/xhtml" not in lowered and "<html" not in html.lower():
        cleaned = " ".join(html.split())
        return cleaned, None

    title = None
    text = ""
    try:
        from readability import Document

        doc = Document(html)
        title = doc.short_title() or doc.title()
        content_html = doc.summary()
        text = _html_to_text(content_html)
    except Exception:
        text = ""
    if not text:
        text = _html_to_text(html)
    cleaned = " ".join(text.split())
    return cleaned, title


# Converts HTML into plain text by removing tags, scripts, and repeated whitespace.
def _html_to_text(html: str) -> str:
    try:
        from bs4 import BeautifulSoup
    except Exception:
        return re.sub(r"<[^>]+>", " ", html)
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text(separator=" ")


# Builds the stored pseudo-document wrapper used for ingested web pages.
def _build_pseudo_doc(title: Optional[str], url: str, crawl_at: str, content_type: str, text: str) -> str:
    header_title = title or url
    lines = [
        f"# {header_title}",
        f"Source: {url}",
        f"Crawled: {crawl_at}",
        f"Content-Type: {content_type or 'unknown'}",
        "",
        text,
    ]
    return "\n".join(lines)


_CAPTION_TIMECODE_RE = re.compile(
    r"^\d{1,2}:\d{2}(?::\d{2})?[.,]\d{3}\s-->\s\d{1,2}:\d{2}(?::\d{2})?[.,]\d{3}"
)


# YouTube caption and transcript helpers.
# Fetches transcript text and metadata for a YouTube video using available captions.
def _fetch_youtube_transcript(

    url: str,
    language: Optional[str],
    allow_auto_captions: Optional[bool],
) -> tuple[str, dict, dict]:
    try:
        from yt_dlp import YoutubeDL
    except Exception as exc:
        raise RuntimeError("yt-dlp is required for YouTube ingestion") from exc

    preferred_language = (language or settings.youtube_default_language or "").strip() or None
    allow_auto = settings.youtube_allow_auto_captions if allow_auto_captions is None else bool(allow_auto_captions)

    ydl_opts = {
        "quiet": True,
        "no_warnings": True,
        "skip_download": True,
        "noplaylist": True,
    }
    with YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=False)
    if not info:
        raise ValueError("Unable to fetch YouTube metadata")

    captions_source, lang, caption_url, caption_ext = _select_caption_track(
        info, preferred_language, allow_auto
    )
    raw = _download_caption_text(caption_url)
    transcript = _parse_caption_text(raw, caption_ext)
    if not transcript:
        raise ValueError("Caption track was empty")

    info_payload = {
        "video_id": info.get("id"),
        "title": info.get("title") or info.get("fulltitle"),
        "channel": info.get("channel") or info.get("uploader"),
        "channel_id": info.get("channel_id") or info.get("uploader_id"),
        "duration_seconds": info.get("duration"),
        "published_at": _format_upload_date(info.get("upload_date")),
    }
    captions_meta = {
        "language": lang,
        "captions_source": captions_source,
        "ext": caption_ext,
    }
    return transcript, info_payload, captions_meta


# Chooses the best matching caption track based on language and caption availability.
def _select_caption_track(
    info: dict,
    preferred_language: Optional[str],
    allow_auto: bool,
) -> tuple[str, str, str, str]:
    """Select the best caption track, prioritising language match over source.

    Priority order:
      1. Manual captions in preferred language
      2. Auto captions in preferred language (if allowed)
      3. Manual captions in English (if preferred wasn't English)
      4. Auto captions in English (if allowed, and preferred wasn't English)
      5. Manual captions in any language
      6. Auto captions in any language (if allowed)
    """
    subtitles = info.get("subtitles") or {}
    auto_caps = (info.get("automatic_captions") or {}) if allow_auto else {}

    preferred_norm = _normalize_language(preferred_language) if preferred_language else None
    preferred_is_english = preferred_norm in (None, "en") or (preferred_norm or "").startswith("en")

    # Tier 1 — preferred language, manual then auto
    selected = _pick_caption_preferred(subtitles, preferred_language)
    if selected:
        lang, url, ext = selected
        return "manual", lang, url, ext
    if auto_caps:
        selected = _pick_caption_preferred(auto_caps, preferred_language)
        if selected:
            lang, url, ext = selected
            return "auto", lang, url, ext

    # Tier 2 — English fallback (skip if preferred was already English)
    if not preferred_is_english:
        selected = _pick_caption_preferred(subtitles, "en")
        if selected:
            lang, url, ext = selected
            return "manual", lang, url, ext
        if auto_caps:
            selected = _pick_caption_preferred(auto_caps, "en")
            if selected:
                lang, url, ext = selected
                return "auto", lang, url, ext

    # Tier 3 — any language, manual then auto
    selected = _pick_caption_any(subtitles)
    if selected:
        lang, url, ext = selected
        return "manual", lang, url, ext
    if auto_caps:
        selected = _pick_caption_any(auto_caps)
        if selected:
            lang, url, ext = selected
            return "auto", lang, url, ext

    if not allow_auto and (info.get("automatic_captions") or {}):
        raise ValueError("No manual captions found. Try enabling auto-captions.")
    raise ValueError("No captions available for this video")


# Finds the best caption track for the preferred language first.
def _pick_caption_preferred(
    captions: dict,
    preferred_language: Optional[str],
) -> Optional[tuple[str, str, str]]:
    """Try to find a caption track matching the preferred language (or English if none set)."""
    if not captions:
        return None
    preferred_norm = _normalize_language(preferred_language) if preferred_language else None
    candidates: list[str] = []
    if preferred_norm:
        candidates.append(preferred_norm)
        if "-" in preferred_norm:
            candidates.append(preferred_norm.split("-", 1)[0])
    else:
        candidates.extend(["en", "en-us", "en-gb"])

    for candidate in candidates:
        for lang_key, formats in captions.items():
            key_norm = _normalize_language(lang_key)
            if key_norm == candidate or key_norm.startswith(candidate):
                selected = _select_caption_format(lang_key, formats)
                if selected:
                    return selected
    return None


# Falls back to the first usable caption track when no preferred match exists.
def _pick_caption_any(
    captions: dict,
) -> Optional[tuple[str, str, str]]:
    """Pick the first available caption track regardless of language."""
    if not captions:
        return None
    for lang_key, formats in captions.items():
        selected = _select_caption_format(lang_key, formats)
        if selected:
            return selected
    return None


# Chooses the best download format for the selected caption track.
def _select_caption_format(lang: str, formats: list[dict]) -> Optional[tuple[str, str, str]]:
    if not formats or not isinstance(formats, list):
        return None
    preferred_exts = ["vtt", "srt", "json3", "srv1", "srv2", "srv3", "ttml"]
    for ext in preferred_exts:
        for item in formats:
            if item.get("ext") == ext and item.get("url"):
                return lang, item["url"], ext
    for item in formats:
        if item.get("url"):
            return lang, item["url"], item.get("ext") or "vtt"
    return None


# Downloads caption text bytes for the selected YouTube caption track.
def _download_caption_text(url: str) -> bytes:
    max_bytes = max(1, settings.youtube_max_bytes)
    retry_statuses = {429, 500, 502, 503, 504}
    max_attempts = 3
    base_delay = 1.5
    last_exc: Optional[Exception] = None
    with httpx.Client(timeout=settings.web_timeout_seconds, follow_redirects=True) as client:
        for attempt in range(1, max_attempts + 1):
            try:
                with client.stream("GET", url) as resp:
                    resp.raise_for_status()
                    total = 0
                    chunks: list[bytes] = []
                    for chunk in resp.iter_bytes():
                        total += len(chunk)
                        if total > max_bytes:
                            raise ValueError("Caption file exceeds size limit")
                        chunks.append(chunk)
                    return b"".join(chunks)
            except httpx.HTTPStatusError as exc:
                last_exc = exc
                status = exc.response.status_code
                if status in retry_statuses and attempt < max_attempts:
                    delay = base_delay * (2 ** (attempt - 1))
                    logger.warning("Caption fetch got %s, retrying in %.1fs", status, delay)
                    time.sleep(delay)
                    continue
                if status == 429:
                    raise ValueError("YouTube rate limit hit; try again later") from exc
                raise
            except httpx.RequestError as exc:
                last_exc = exc
                if attempt < max_attempts:
                    delay = base_delay * (2 ** (attempt - 1))
                    logger.warning("Caption fetch failed, retrying in %.1fs: %s", delay, exc)
                    time.sleep(delay)
                    continue
                raise
    if last_exc:
        raise last_exc
    raise RuntimeError("Failed to download captions")


# Parses downloaded caption content into plain transcript text.
def _parse_caption_text(raw: bytes, ext: str) -> str:
    text = raw.decode("utf-8", errors="ignore")
    ext_lower = (ext or "").lower()
    if ext_lower in {"vtt", "srt"}:
        return _strip_vtt_srt(text)
    if ext_lower in {"srv1", "srv2", "srv3", "ttml", "xml"}:
        return _strip_xml(text)
    if ext_lower == "json3":
        return _parse_json3(text)
    cleaned = _strip_vtt_srt(text)
    return cleaned or _strip_xml(text)


# Removes timestamps and markup from VTT or SRT caption text.
def _strip_vtt_srt(text: str) -> str:
    lines: list[str] = []
    for line in text.splitlines():
        cleaned = line.strip()
        if not cleaned:
            continue
        if cleaned.startswith("WEBVTT") or cleaned.startswith("NOTE"):
            continue
        if cleaned.isdigit():
            continue
        if _CAPTION_TIMECODE_RE.match(cleaned):
            continue
        cleaned = re.sub(r"<[^>]+>", "", cleaned)
        lines.append(cleaned)
    return " ".join(lines).strip()


# Removes XML tags from caption payloads that still contain markup.
def _strip_xml(text: str) -> str:
    cleaned = re.sub(r"<[^>]+>", " ", text)
    return " ".join(cleaned.split()).strip()


# Extracts transcript text from YouTube's JSON3 caption format.
def _parse_json3(text: str) -> str:
    try:
        payload = json.loads(text)
    except json.JSONDecodeError:
        return ""
    parts: list[str] = []
    for event in payload.get("events", []) or []:
        for seg in event.get("segs", []) or []:
            seg_text = seg.get("utf8")
            if seg_text:
                parts.append(seg_text)
    return " ".join(parts).strip()


# Normalizes language codes into a simple lowercase value.
def _normalize_language(value: Optional[str]) -> str:
    if not value:
        return ""
    return value.strip().lower().replace("_", "-")


# Formats a compact YouTube upload date into an ISO-like calendar date.
def _format_upload_date(value: Optional[str]) -> Optional[str]:
    if not value or len(value) != 8:
        return None
    return f"{value[:4]}-{value[4:6]}-{value[6:]}"


# Formats a duration in seconds into a human-readable string.
def _format_duration(seconds: Optional[int]) -> Optional[str]:
    if seconds is None:
        return None
    try:
        total = int(seconds)
    except (TypeError, ValueError):
        return None
    hours = total // 3600
    minutes = (total % 3600) // 60
    secs = total % 60
    if hours:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


# Builds the stored pseudo-document wrapper used for ingested YouTube transcripts.
def _build_youtube_pseudo_doc(info: dict, url: str, fetched_at: str, captions_meta: dict, text: str) -> str:
    title = info.get("title") or "YouTube Video"
    lines: list[str] = [f"# {title}"]
    if url:
        lines.append(f"Source: {url}")
    if info.get("video_id"):
        lines.append(f"Video ID: {info['video_id']}")
    if info.get("channel"):
        lines.append(f"Channel: {info['channel']}")
    if info.get("published_at"):
        lines.append(f"Published: {info['published_at']}")
    duration = _format_duration(info.get("duration_seconds"))
    if duration:
        lines.append(f"Duration: {duration}")
    if captions_meta.get("language"):
        lines.append(f"Language: {captions_meta['language']}")
    if captions_meta.get("captions_source"):
        lines.append(f"Captions: {captions_meta['captions_source']}")
    lines.append(f"Fetched: {fetched_at}")
    lines.append("")
    lines.append(text)
    return "\n".join(lines)


# Uploads the generated pseudo-document back into Supabase Storage.
def _upload_pseudo_doc(client: Client, storage_path: str, content: str) -> None:
    if not storage_path:
        raise ValueError("Storage path missing for pseudo document")
    payload = content.encode("utf-8")
    try:
        client.storage.from_(settings.storage_bucket).remove([storage_path])
    except Exception:
        pass
    client.storage.from_(settings.storage_bucket).upload(
        storage_path,
        payload,
        {"content-type": "text/markdown"},
    )


# Text extraction, chunking, and embedding helpers.
# Routes raw file bytes to the correct text extractor based on file extension.
def _extract_text(raw: bytes, storage_path: str) -> str:

    # Route by file extension to the right extractor.
    ext = Path(storage_path).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(raw)
    if ext == ".docx":
        return _extract_docx(raw)
    if ext in {".md", ".txt"}:
        return raw.decode("utf-8", errors="ignore")
    return raw.decode("utf-8", errors="ignore")


# Extracts readable text from a PDF file.
def _extract_pdf(raw: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(raw))
    except Exception as exc:
        raise ValueError(f"Could not read PDF: {exc}") from exc
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n".join(pages)


# Extracts readable text from a DOCX file.
def _extract_docx(raw: bytes) -> str:
    import docx  # local import to avoid unused dependency warnings if not used

    doc = docx.Document(io.BytesIO(raw))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


# Normalizes extracted text into a cleaner single-space format.
def _normalize_text(text: str) -> str:
    return " ".join(text.replace("\r", "\n").split())


# Trims text to a maximum length while preserving a readable ending.
def _trim_text(text: str, max_chars: int) -> str:
    cleaned = _normalize_text(text)
    if len(cleaned) <= max_chars:
        return cleaned
    return f"{cleaned[:max_chars].rstrip()}..."


# Chunks text, embeds it, stores the vectors, and triggers reminder detection for a source.
def _ingest_text_for_source(
    client: Client,
    source_id: str,
    hub_id: str,
    text: str,
    extra_metadata: Optional[dict],
) -> int:
    chunks = _chunk_text(text, settings.chunk_size, settings.chunk_overlap)
    if not chunks:
        raise ValueError("No chunks produced from extracted text")
    if not _source_exists(client, source_id):
        logger.info("Source %s deleted before ingest; skipping.", source_id)
        return 0
    ingest_started_at = datetime.now(timezone.utc)
    ingest_timestamp = ingest_started_at.isoformat()
    embeddings = _embed_chunks(chunks)
    if not _source_exists(client, source_id):
        logger.info("Source %s deleted during embed; skipping insert.", source_id)
        return 0
    _insert_chunks(client, source_id, hub_id, chunks, embeddings, ingest_timestamp)
    _clear_existing_chunks_before(client, source_id, ingest_timestamp)
    existing_metadata = _get_source_metadata(client, source_id)
    metadata = {
        "chunk_count": len(chunks),
        "embedding_model": settings.embedding_model,
        "chunk_size": settings.chunk_size,
        "chunk_overlap": settings.chunk_overlap,
    }
    merged = {**existing_metadata, **metadata}
    if extra_metadata:
        merged.update(extra_metadata)
    _update_source(client, source_id, status="complete", ingestion_metadata=merged)
    try:
        _detect_and_store_reminders(client, source_id, hub_id, text)
    except Exception:
        logger.exception("Reminder detection failed for source %s", source_id)
    return len(chunks)


# Splits text into overlapping word chunks for embedding.
def _chunk_text(text: str, chunk_size: int, overlap: int) -> List[str]:
    words = text.split()
    if not words:
        return []
    chunks: List[str] = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk_words = words[start:end]
        chunks.append(" ".join(chunk_words))
        if end == len(words):
            break
        start = max(end - overlap, 0)
    return chunks


# Requests embeddings for the prepared text chunks from OpenAI.
def _embed_chunks(chunks: List[str]) -> List[List[float]]:
    if not settings.openai_api_key:
        raise RuntimeError("OPENAI_API_KEY missing in worker environment")
    client = OpenAI(api_key=settings.openai_api_key)
    embeddings: List[List[float]] = []
    for batch in _batch(chunks, 64):
        # Batch requests to keep payload size under API limits.
        response = client.embeddings.create(model=settings.embedding_model, input=batch)
        embeddings.extend([item.embedding for item in response.data])
    return embeddings


# Writes embedded chunk records into the database in batches.
def _insert_chunks(
    client: Client,
    source_id: str,
    hub_id: str,
    chunks: List[str],
    embeddings: List[List[float]],
    created_at: str,
) -> None:
    rows = []
    for idx, (text, embedding) in enumerate(zip(chunks, embeddings, strict=False)):
        rows.append(
            {
                "source_id": source_id,
                "hub_id": hub_id,
                "chunk_index": idx,
                "text": text,
                "embedding": embedding,
                "token_count": len(text.split()),
                "metadata": {"word_count": len(text.split())},
                "created_at": created_at,
            }
        )
    for batch in _batch(rows, 100):
        client.table("source_chunks").insert(batch).execute()


# Deletes older chunk rows for the same source before new ones are inserted.
def _clear_existing_chunks_before(client: Client, source_id: str, cutoff: str) -> None:
    client.table("source_chunks").delete().eq("source_id", source_id).lt("created_at", cutoff).execute()


# Checks whether the source row still exists before continuing ingestion work.
def _source_exists(client: Client, source_id: str) -> bool:
    response = client.table("sources").select("id").eq("id", source_id).limit(1).execute()
    return bool(response.data)


# Loads source metadata fields that are needed during reminder detection.
def _get_source_metadata(client: Client, source_id: str) -> dict:
    response = (
        client.table("sources")
        .select("ingestion_metadata")
        .eq("id", source_id)
        .limit(1)
        .execute()
    )
    if not response.data:
        return {}
    metadata = response.data[0].get("ingestion_metadata") or {}
    return metadata if isinstance(metadata, dict) else {}


# Updates the source status and related bookkeeping fields in the database.
def _update_source(
    client: Client,
    source_id: str,
    status: str,
    failure_reason: Optional[str] = None,
    ingestion_metadata: Optional[dict] = None,
    clear_failure_reason: bool = False,
) -> None:
    payload: dict = {"status": status}
    if failure_reason is not None or clear_failure_reason:
        payload["failure_reason"] = failure_reason
    if ingestion_metadata is not None:
        payload["ingestion_metadata"] = ingestion_metadata
    client.table("sources").update(payload).eq("id", source_id).execute()


# Yields items in fixed-size batches for APIs that expect chunked writes.
def _batch(items: List, size: int) -> Iterable[List]:
    for i in range(0, len(items), size):
        yield items[i : i + size]

# Reminder detection pipeline (regex + spaCy) and dispatch.

MAX_TEXT_CHARS = 200_000
MIN_CONFIDENCE = 0.7
MAX_CANDIDATES = 6
DATE_KEYWORDS = (
    "due",
    "deadline",
    "submit",
    "submission",
    "by",
    "before",
    "no later than",
    "must be received",
    "final date",
    "window",
)
DATE_TIME_RE = re.compile(r"\b(\d{1,2}:\d{2}\b|\d{1,2}\s*(am|pm)\b)", re.IGNORECASE)
TIME_ONLY_RE = re.compile(r"^\s*\d{1,2}(:\d{2})?\s*(am|pm)?\s*$", re.IGNORECASE)
SENTENCE_BOUNDARY_RE = re.compile(r"[.!?]")
MONTH_PATTERN = (
    r"jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|"
    r"aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?"
)
RANGE_NUMERIC_RE = re.compile(
    r"(?P<start>\d{1,2}[/-]\d{1,2}(?:[/-]\d{2,4})?)\s*(?:-|\u2013|\u2014)\s*"
    r"(?P<end>\d{1,2}[/-]\d{1,2}(?:[/-]\d{2,4})?)"
)
RANGE_MONTH_RE = re.compile(
    rf"(?P<start>\d{{1,2}}(?:st|nd|rd|th)?)\s*(?:-|\u2013|\u2014)\s*"
    rf"(?P<end>\d{{1,2}}(?:st|nd|rd|th)?)\s+(?P<month>{MONTH_PATTERN})(?:\s+(?P<year>\d{{4}}))?",
    re.IGNORECASE,
)
DATE_REGEXES = [
    re.compile(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b"),
    re.compile(r"\b\d{4}-\d{1,2}-\d{1,2}\b"),
    re.compile(rf"\b(?:{MONTH_PATTERN})\s+\d{{1,2}}(?:st|nd|rd|th)?(?:,?\s+\d{{4}})?\b", re.IGNORECASE),
    re.compile(rf"\b\d{{1,2}}(?:st|nd|rd|th)?\s+(?:{MONTH_PATTERN})(?:\s+\d{{4}})?\b", re.IGNORECASE),
]

_NLP = None


# Reminder detection and date parsing helpers.
# Finds possible reminder dates in source text and stores any valid reminder candidates.
def _detect_and_store_reminders(client: Client, source_id: str, hub_id: str, text: str) -> None:

    # Cap text length for deterministic runtime; candidates are deduped via upsert.
    cleaned = text[:MAX_TEXT_CHARS]
    candidates = _find_date_candidates(cleaned, settings.default_timezone)
    if not candidates:
        return
    rows = []
    for candidate in candidates:
        rows.append(
            {
                "hub_id": hub_id,
                "source_id": source_id,
                "detected_by": candidate["detected_by"],
                "snippet": candidate["snippet"],
                "snippet_hash": candidate["snippet_hash"],
                "due_at": candidate["due_at"],
                "timezone": candidate["timezone"],
                "title_suggestion": candidate["title_suggestion"],
                "confidence": candidate["confidence"],
                "status": "pending",
            }
        )
    for batch in _batch(rows, 50):
        client.table("reminder_candidates").upsert(
            batch, on_conflict="source_id,due_at,snippet_hash"
        ).execute()


# Builds reminder date candidates from multiple date-detection strategies.
def _find_date_candidates(text: str, timezone_name: str) -> List[dict]:
    mentions = _collect_date_mentions(text)
    now = datetime.now(timezone.utc)
    candidates: List[dict] = []
    seen_keys: set[tuple[str, str]] = set()
    for mention in mentions:
        date_text = mention["text"]
        if re.fullmatch(r"\d{4}", date_text.strip()):
            continue
        range_end = _extract_range_end(date_text)
        if range_end:
            date_text = range_end
        if _looks_historical_or_vague_date(date_text):
            continue
        if mention["method"] == "ner" and _is_numeric_only(date_text):
            continue
        if _is_day_only(date_text):
            continue
        if _is_time_only(date_text):
            continue
        if _is_week_reference(date_text):
            continue
        time_hint = _extract_time_hint(text, mention["start"], mention["end"])
        parse_text = date_text
        # Attach a nearby time to otherwise date-only mentions when one is available.
        if time_hint and not _has_time(date_text):
            parse_text = f"{date_text} {time_hint}"
        parsed = _parse_date_text(parse_text, timezone_name, now)
        if not parsed:
            continue
        if not _is_reasonable_date(parsed, now):
            continue
        snippet = _extract_snippet(text, mention["start"], mention["end"])
        has_keyword = _has_keyword(snippet) or _has_keyword_near(text, mention["start"], mention["end"])
        if _looks_relative(date_text) and not has_keyword:
            continue
        if _is_repeated_date(text, date_text) and not has_keyword:
            continue
        if _is_numeric_date(parse_text) and not has_keyword:
            continue
        confidence = _score_candidate(mention["method"], snippet, parse_text)
        if confidence < MIN_CONFIDENCE:
            continue
        snippet_hash = _hash_snippet(snippet)
        # Deduplicate by parsed due date plus snippet so repeated mentions do not create duplicate candidates.
        key = (parsed.isoformat(), snippet_hash)
        if key in seen_keys:
            continue
        seen_keys.add(key)
        candidates.append(
            {
                "detected_by": mention["method"],
                "snippet": snippet,
                "snippet_hash": snippet_hash,
                "due_at": parsed.isoformat(),
                "timezone": timezone_name,
                "title_suggestion": _build_title(snippet),
                "confidence": confidence,
            }
        )
        if len(candidates) >= MAX_CANDIDATES:
            break
    return _dedupe_best_candidates(candidates)


# Collects standalone date mentions from the source text.
def _collect_date_mentions(text: str) -> List[dict]:
    mentions: List[dict] = []
    mentions.extend(_collect_range_mentions(text))
    for regex in DATE_REGEXES:
        for match in regex.finditer(text):
            mentions.append(
                {"text": match.group(0), "start": match.start(), "end": match.end(), "method": "regex"}
            )
    nlp = _get_nlp()
    if nlp is None:
        return mentions
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ != "DATE":
            continue
        mentions.append(
            {"text": ent.text, "start": ent.start_char, "end": ent.end_char, "method": "ner"}
        )
    return mentions


# Deduplicates reminder candidates and keeps the strongest match for each date.
def _dedupe_best_candidates(candidates: List[dict]) -> List[dict]:
    best_by_snippet: dict[str, dict] = {}
    for candidate in candidates:
        key = candidate.get("snippet_hash") or ""
        best = best_by_snippet.get(key)
        if best is None:
            best_by_snippet[key] = candidate
            continue
        if candidate["confidence"] > best["confidence"]:
            best_by_snippet[key] = candidate
        elif candidate["confidence"] == best["confidence"] and candidate["due_at"] < best["due_at"]:
            best_by_snippet[key] = candidate
    filtered: List[dict] = []
    seen: set[str] = set()
    for candidate in candidates:
        key = candidate.get("snippet_hash") or ""
        if key in seen:
            continue
        if best_by_snippet.get(key) is candidate:
            filtered.append(candidate)
            seen.add(key)
    return filtered


# Collects date range mentions so start and end dates can be interpreted together.
def _collect_range_mentions(text: str) -> List[dict]:
    mentions: List[dict] = []
    for match in RANGE_NUMERIC_RE.finditer(text):
        end_text = _normalize_numeric_range_end(match.group("start"), match.group("end"))
        if not end_text:
            continue
        mentions.append(
            {"text": end_text, "start": match.start("end"), "end": match.end("end"), "method": "range"}
        )
    for match in RANGE_MONTH_RE.finditer(text):
        end_text = f"{match.group('end')} {match.group('month')}"
        if match.group("year"):
            end_text = f"{end_text} {match.group('year')}"
        mentions.append(
            {"text": end_text, "start": match.start("end"), "end": match.end("end"), "method": "range"}
        )
    return mentions


# Lazily loads and caches the spaCy model used for date extraction support.
def _get_nlp():
    global _NLP
    if _NLP is not None:
        return _NLP
    try:
        # Soft-fail if spaCy or the model is unavailable.
        _NLP = spacy.load("en_core_web_sm")
    except Exception:
        _NLP = None
    return _NLP


# Parses a detected date phrase into a timezone-aware datetime.
def _parse_date_text(date_text: str, timezone_name: str, now: datetime) -> Optional[datetime]:
    iso_match = re.fullmatch(r"\s*(\d{4}-\d{2}-\d{2})(?:[ T](\d{1,2}:\d{2})(?::\d{2})?)?\s*", date_text)
    if iso_match:
        base = iso_match.group(1)
        time_part = iso_match.group(2)
        parsed = datetime.fromisoformat(f"{base} {time_part}" if time_part else base)
        if parsed.tzinfo is None:
            tz = _safe_zoneinfo(timezone_name) or timezone.utc
            parsed = parsed.replace(tzinfo=tz)
        if not time_part:
            parsed = parsed.replace(hour=9, minute=0, second=0, microsecond=0)
        return parsed.astimezone(timezone.utc)
    settings_payload = {
        "PREFER_DATES_FROM": "future",
        "RELATIVE_BASE": now,
        "RETURN_AS_TIMEZONE_AWARE": True,
        "TIMEZONE": timezone_name,
        "TO_TIMEZONE": "UTC",
        "DATE_ORDER": "DMY",
    }
    parsed = dateparser.parse(date_text, settings=settings_payload)
    if not parsed:
        return None
    if parsed.tzinfo is None:
        tz = _safe_zoneinfo(timezone_name) or timezone.utc
        parsed = parsed.replace(tzinfo=tz)
    if not _has_time(date_text):
        # Default date-only mentions to 9:00 local time.
        parsed = parsed.replace(hour=9, minute=0, second=0, microsecond=0)
    return parsed.astimezone(timezone.utc)


# Safely resolves a timezone name into a `ZoneInfo` object.
def _safe_zoneinfo(name: str) -> Optional[ZoneInfo]:
    try:
        return ZoneInfo(name)
    except Exception:
        return None


# Extracts a local text snippet around a detected date for context and scoring.
def _extract_snippet(text: str, start: int, end: int, radius: int = 120) -> str:
    snippet_start = max(0, start - radius)
    snippet_end = min(len(text), end + radius)
    window = text[snippet_start:snippet_end]
    if not window:
        return ""
    local_start = max(0, start - snippet_start)
    local_end = max(0, end - snippet_start)
    sentence_start = _find_sentence_start(window, local_start)
    sentence_end = _find_sentence_end(window, local_end)
    if sentence_start >= sentence_end:
        sentence_start = max(0, local_start - radius // 2)
        sentence_end = min(len(window), local_end + radius // 2)
    snippet = window[sentence_start:sentence_end]
    snippet = re.sub(r"\s+", " ", snippet).strip()
    snippet = re.sub(r"\b\d{1,2}\)\s*", "", snippet)
    return snippet[:280]


# Builds a short reminder title from the surrounding source snippet.
def _build_title(snippet: str) -> str:
    cleaned = snippet.strip()
    if len(cleaned) <= 80:
        return cleaned
    return f"{cleaned[:77].rstrip()}..."


# Hashes a snippet so repeated reminder candidates can be deduplicated.
def _hash_snippet(snippet: str) -> str:
    return hashlib.sha256(snippet.lower().encode("utf-8")).hexdigest()


# Scores a reminder candidate so better matches can win during deduplication.
def _score_candidate(method: str, snippet: str, date_text: str) -> float:
    score = 0.3
    if method == "regex":
        score += 0.35
    if method == "range":
        score += 0.35
    if method == "ner":
        score += 0.25
    if _has_keyword(snippet):
        score += 0.2
    if _has_time(date_text):
        score += 0.1
    if _is_ambiguous_numeric(date_text):
        score -= 0.1
    if _looks_mathy(snippet):
        score -= 0.2
    return max(0.0, min(0.95, score))


# Checks whether a snippet contains reminder-related keywords.
def _has_keyword(snippet: str) -> bool:
    lowered = snippet.lower()
    return any(keyword in lowered for keyword in DATE_KEYWORDS)


# Checks whether a date phrase also contains an explicit time.
def _has_time(text: str) -> bool:
    return bool(DATE_TIME_RE.search(text))


# Checks whether reminder keywords appear near a detected date mention.
def _has_keyword_near(text: str, start: int, end: int, window: int = 120) -> bool:
    if start < 0 or end < 0:
        return False
    win_start = max(0, start - window)
    win_end = min(len(text), end + window)
    return _has_keyword(text[win_start:win_end])


# Checks whether a date phrase is relative rather than a concrete scheduled date.
def _looks_relative(date_text: str) -> bool:
    value = date_text.strip().lower()
    return bool(
        re.search(
            r"\b(next|tomorrow|today|tonight|this|within|in\s+\d+|end of|end-of|after)\b",
            value,
        )
    )


# Checks whether a parsed phrase looks like a time without a date.
def _is_time_only(date_text: str) -> bool:
    return bool(TIME_ONLY_RE.match(date_text.strip()))


# Checks whether a phrase refers to a week rather than a specific event date.
def _is_week_reference(date_text: str) -> bool:
    return bool(re.search(r"\b(?:week|wk)\s*\d{1,2}\b", date_text.strip().lower()))


# Checks whether a phrase names only a weekday without enough scheduling detail.
def _is_day_only(date_text: str) -> bool:
    return bool(re.fullmatch(r"\d{1,2}(st|nd|rd|th)?", date_text.strip(), re.IGNORECASE))


# Checks whether a phrase resembles a numeric calendar date.
def _is_numeric_date(date_text: str) -> bool:
    value = date_text.strip()
    if re.search(r"[a-zA-Z]", value):
        return False
    return bool(re.search(r"[/-]", value)) or value.isdigit()


# Checks whether a phrase is only digits without enough date context.
def _is_numeric_only(date_text: str) -> bool:
    return bool(re.fullmatch(r"\d+", date_text.strip()))


# Extracts the trailing date text from a detected date range.
def _extract_range_end(date_text: str) -> Optional[str]:
    match = RANGE_NUMERIC_RE.search(date_text)
    if match:
        return _normalize_numeric_range_end(match.group("start"), match.group("end"))
    match = RANGE_MONTH_RE.search(date_text)
    if match:
        end_text = f"{match.group('end')} {match.group('month')}"
        if match.group("year"):
            end_text = f"{end_text} {match.group('year')}"
        return end_text
    return None


# Normalizes short range endings so they can be parsed with the range start.
def _normalize_numeric_range_end(start_text: str, end_text: str) -> Optional[str]:
    start_parts = re.split(r"[/-]", start_text)
    end_parts = re.split(r"[/-]", end_text)
    if len(end_parts) == 2 and len(start_parts) >= 3:
        return f"{end_parts[0]}/{end_parts[1]}/{start_parts[2]}"
    if len(end_parts) >= 2:
        return "/".join(end_parts)
    return None


# Checks whether a date phrase appears too many times to be a useful reminder anchor.
def _is_repeated_date(text: str, date_text: str) -> bool:
    needle = date_text.strip().lower()
    if len(needle) < 4:
        return False
    return text.lower().count(needle) >= 3


# Rejects dates that look historical, vague, or otherwise unsuitable for reminders.
def _looks_historical_or_vague_date(date_text: str) -> bool:
    value = date_text.strip().lower()
    if re.search(r"\b\d{3,4}s\b", value):
        return True
    if re.fullmatch(r"(?:in|by|during|around|circa|c\.|approx\.?|about)?\s*\d{4}", value):
        return True
    if re.search(r"\b\d{1,2}(st|nd|rd|th)\s+century\b", value):
        return True
    if re.search(r"\b(?:bc|bce|ad|ce)\b", value):
        return True
    return False


# Rejects numeric date phrases that are too ambiguous to trust.
def _is_ambiguous_numeric(date_text: str) -> bool:
    match = re.match(r"\b(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})\b", date_text.strip())
    if not match:
        return False
    first = int(match.group(1))
    second = int(match.group(2))
    return first <= 12 and second <= 12 and first != second


# Finds the start of the local sentence around a detected match.
def _find_sentence_start(window: str, idx: int) -> int:
    start = 0
    for match in SENTENCE_BOUNDARY_RE.finditer(window[:idx]):
        start = match.end()
    while start < len(window) and window[start].isspace():
        start += 1
    return start


# Finds the end of the local sentence around a detected match.
def _find_sentence_end(window: str, idx: int) -> int:
    match = SENTENCE_BOUNDARY_RE.search(window[idx:])
    if match:
        end = idx + match.end()
    else:
        end = len(window)
    while end > 0 and end < len(window) and window[end - 1].isspace():
        end -= 1
    return end


# Extracts a nearby time expression that may refine the parsed reminder date.
def _extract_time_hint(text: str, start: int, end: int, window: int = 60) -> Optional[str]:
    if start < 0 or end < 0:
        return None
    win_start = max(0, start - window)
    win_end = min(len(text), end + window)
    window_text = text[win_start:win_end]
    local_start = max(0, start - win_start)
    local_end = max(0, end - win_start)
    sentence_start = _find_sentence_start(window_text, local_start)
    sentence_end = _find_sentence_end(window_text, local_end)
    sentence_text = window_text[sentence_start:sentence_end]
    matches = list(DATE_TIME_RE.finditer(sentence_text))
    if not matches:
        return None
    mention_center = (start + end) / 2
    best = None
    best_distance = None
    for match in matches:
        match_center = win_start + sentence_start + (match.start() + match.end()) / 2
        distance = abs(match_center - mention_center)
        if best_distance is None or distance < best_distance:
            best_distance = distance
            best = match.group(0)
    return best


# Checks whether a parsed reminder date falls within a sensible time range.
def _is_reasonable_date(value: datetime, now: datetime) -> bool:
    if value < now - timedelta(days=30):
        return False
    if value > now + timedelta(days=365 * 2):
        return False
    return True


# Rejects snippets that look more like formulas than natural-language reminders.
def _looks_mathy(snippet: str) -> bool:
    if not snippet:
        return False
    letters = sum(ch.isalpha() for ch in snippet)
    digits = sum(ch.isdigit() for ch in snippet)
    symbols = sum(ch in "=+-*/^_" for ch in snippet)
    if letters == 0:
        return digits > 0
    digit_ratio = digits / max(letters, 1)
    symbol_ratio = symbols / max(len(snippet), 1)
    lowered = snippet.lower()
    if "mod " in lowered or "modulo" in lowered:
        return True
    return digit_ratio >= 0.6 or symbol_ratio > 0.05


_SUGGESTION_YOUTUBE_ID_RE = re.compile(r"^[A-Za-z0-9_-]{11}$")


# Source suggestion scanning and discovery helpers.
# Builds the Redis client used for worker locks and background coordination.
def _get_redis_client() -> redis.Redis:

    redis_url = settings.redis_url
    client_kwargs: dict = {}

    if redis_url.lower().startswith("rediss://"):
        parsed = urlparse(redis_url)
        query = parse_qs(parsed.query, keep_blank_values=True)
        ssl_cert_reqs_raw = (query.get("ssl_cert_reqs") or [None])[-1]
        if isinstance(ssl_cert_reqs_raw, str):
            mapped = {
                "cert_none": ssl.CERT_NONE,
                "none": ssl.CERT_NONE,
                "0": ssl.CERT_NONE,
                "cert_optional": ssl.CERT_OPTIONAL,
                "optional": ssl.CERT_OPTIONAL,
                "1": ssl.CERT_OPTIONAL,
                "cert_required": ssl.CERT_REQUIRED,
                "required": ssl.CERT_REQUIRED,
                "2": ssl.CERT_REQUIRED,
            }.get(ssl_cert_reqs_raw.strip().lower())
            if mapped is not None:
                client_kwargs["ssl_cert_reqs"] = mapped
                if mapped == ssl.CERT_NONE:
                    client_kwargs["ssl_check_hostname"] = False
                query.pop("ssl_cert_reqs", None)
                redis_url = urlunparse(
                    (
                        parsed.scheme,
                        parsed.netloc,
                        parsed.path,
                        parsed.params,
                        urlencode(query, doseq=True),
                        parsed.fragment,
                    )
                )

    return redis.Redis.from_url(redis_url, **client_kwargs)


# Claims a per-hub Redis lock so only one worker scans suggestions at a time.
def _acquire_source_suggestion_lock(hub_id: str) -> Optional[tuple[redis.Redis, str, str]]:
    client = _get_redis_client()
    key = f"locks:source-suggestions:{hub_id}"
    token = str(uuid.uuid4())
    try:
        acquired = client.set(key, token, nx=True, ex=max(60, settings.suggested_sources_lock_ttl_seconds))
    except Exception:
        client.close()
        raise
    if not acquired:
        client.close()
        return None
    return client, key, token


# Releases a previously acquired source-suggestion Redis lock.
def _release_source_suggestion_lock(lock: tuple[redis.Redis, str, str]) -> None:
    client, key, token = lock
    try:
        current = client.get(key)
        if current is not None and current.decode("utf-8", errors="ignore") == token:
            client.delete(key)
    finally:
        client.close()


# Loads hubs that are eligible for a new source-suggestion scan.
def _list_eligible_source_suggestion_hubs(client: Client, now: Optional[datetime] = None) -> list[dict]:
    now = now or datetime.now(timezone.utc)
    hubs_response = client.table("hubs").select("id,last_source_suggestion_scan_at").execute()
    hubs = hubs_response.data or []
    if not hubs:
        return []

    complete_counts: dict[str, int] = defaultdict(int)
    for row in client.table("sources").select("hub_id").eq("status", "complete").execute().data or []:
        hub_id = str(row.get("hub_id") or "")
        if hub_id:
            complete_counts[hub_id] += 1

    active_cutoff = now - timedelta(days=max(1, settings.suggested_sources_active_days))
    active_hub_ids = {
        str(row.get("hub_id"))
        for row in (
            client.table("hub_members")
            .select("hub_id")
            .not_.is_("accepted_at", "null")
            .gte("last_accessed_at", active_cutoff.isoformat())
            .execute()
            .data
            or []
        )
        if row.get("hub_id")
    }
    pending_hub_ids = {
        str(row.get("hub_id"))
        for row in client.table("source_suggestions").select("hub_id").eq("status", "pending").execute().data or []
        if row.get("hub_id")
    }
    return _filter_eligible_source_suggestion_hubs(
        hubs,
        complete_source_counts=complete_counts,
        active_hub_ids=active_hub_ids,
        pending_hub_ids=pending_hub_ids,
        now=now,
    )


# Filters hubs down to those that are active, ready, and outside cooldown windows.
def _filter_eligible_source_suggestion_hubs(
    hubs: list[dict],
    *,
    complete_source_counts: dict[str, int],
    active_hub_ids: set[str],
    pending_hub_ids: set[str],
    now: datetime,
) -> list[dict]:
    eligible: list[dict] = []
    cooldown = timedelta(minutes=max(1, settings.suggested_sources_hub_cooldown_minutes))
    min_sources = max(1, settings.suggested_sources_min_complete_sources)

    for hub in hubs:
        hub_id = str(hub.get("id") or "")
        if not hub_id:
            continue
        if hub_id in pending_hub_ids:
            continue
        if complete_source_counts.get(hub_id, 0) < min_sources:
            continue
        if hub_id not in active_hub_ids:
            continue
        last_scan = _parse_iso(hub.get("last_source_suggestion_scan_at"))
        if last_scan is not None and last_scan > now - cooldown:
            continue
        eligible.append(hub)
    return eligible


# Generates and stores new source suggestions for a single eligible hub.
def _generate_source_suggestions_for_hub(client: Client, hub_id: str, now: Optional[datetime] = None) -> dict:
    now = now or datetime.now(timezone.utc)
    seed_source_ids, context_text = _build_source_suggestion_context(client, hub_id)
    if len(seed_source_ids) < settings.suggested_sources_min_complete_sources or not context_text.strip():
        _mark_source_suggestion_scan(client, hub_id, now=now, generated=False)
        return {"inserted": 0, "reason": "insufficient_context"}

    discovered, search_metadata = _discover_source_suggestions(context_text)
    normalized = _normalize_source_suggestion_candidates(
        discovered,
        hub_id=hub_id,
        seed_source_ids=seed_source_ids,
        search_metadata=search_metadata,
    )
    existing_suggestion_targets = _load_existing_source_suggestion_targets(client, hub_id)
    existing_source_targets = _load_existing_source_targets(client, hub_id)
    pending_rows = _filter_new_source_suggestions(
        normalized,
        existing_source_targets=existing_source_targets,
        existing_suggestion_targets=existing_suggestion_targets,
        limit=max(1, settings.suggested_sources_batch_limit),
    )
    if pending_rows:
        client.table("source_suggestions").insert(pending_rows).execute()
    _mark_source_suggestion_scan(client, hub_id, now=now, generated=bool(pending_rows))
    return {"inserted": len(pending_rows)}


# Builds the LLM context bundle used to discover new source suggestions.
def _build_source_suggestion_context(client: Client, hub_id: str) -> tuple[list[str], str]:
    source_rows = (
        client.table("sources")
        .select("id,type,original_name,ingestion_metadata,created_at")
        .eq("hub_id", hub_id)
        .eq("status", "complete")
        .order("created_at", desc=True)
        .limit(max(1, settings.suggested_sources_context_limit))
        .execute()
        .data
        or []
    )
    if not source_rows:
        return [], ""

    source_ids = [str(row.get("id")) for row in source_rows if row.get("id")]
    chunk_rows = []
    if source_ids:
        chunk_rows = (
            client.table("source_chunks")
            .select("source_id,chunk_index,text")
            .in_("source_id", source_ids)
            .order("chunk_index")
            .execute()
            .data
            or []
        )

    excerpts: dict[str, list[str]] = defaultdict(list)
    max_chunks = max(1, settings.suggested_sources_chunks_per_source)
    for row in chunk_rows:
        source_id = str(row.get("source_id") or "")
        text = _normalize_text(str(row.get("text") or ""))
        if not source_id or not text or len(excerpts[source_id]) >= max_chunks:
            continue
        excerpts[source_id].append(_trim_text(text, 500))

    blocks: list[str] = []
    for row in source_rows:
        source_id = str(row.get("id") or "")
        source_type = str(row.get("type") or "file")
        title = str(row.get("original_name") or source_id)
        metadata = row.get("ingestion_metadata") if isinstance(row.get("ingestion_metadata"), dict) else {}
        lines = [f"Source ID: {source_id}", f"Type: {source_type}", f"Title: {title}"]
        if source_type == "web":
            url = metadata.get("final_url") or metadata.get("url")
            if url:
                lines.append(f"URL: {url}")
        if source_type == "youtube":
            channel = metadata.get("channel")
            if channel:
                lines.append(f"Channel: {channel}")
            video_id = metadata.get("video_id")
            if video_id:
                lines.append(f"Video ID: {video_id}")
        for idx, excerpt in enumerate(excerpts.get(source_id, []), start=1):
            lines.append(f"Excerpt {idx}: {excerpt}")
        blocks.append("\n".join(lines))
    return source_ids, "\n\n".join(blocks)


# Calls the LLM and web search tools to discover candidate source suggestions.
def _discover_source_suggestions(context_text: str) -> tuple[list[dict], dict]:
    if not settings.openai_api_key:
        logger.warning("Skipping source suggestions because OPENAI_API_KEY is missing")
        return [], {"error": "missing_openai_api_key"}

    llm_client = OpenAI(api_key=settings.openai_api_key)
    responses_client = getattr(llm_client, "responses", None)
    if responses_client is None:
        logger.warning("Skipping source suggestions because Responses API is unavailable")
        return [], {"error": "responses_api_unavailable"}

    system_prompt = (
        "You find suggested sources for a study hub. Use web search to find public resources that complement the existing hub material. "
        "Return only a JSON array with up to 6 objects. Each object must include: "
        "type ('web' or 'youtube'), url, title, description, rationale, confidence. "
        "Prefer authoritative pages and relevant YouTube videos. When relevant, include at least 1 YouTube video in the returned set. "
        "Avoid login pages, homepages with no clear relevance, PDFs, and duplicates."
    )
    user_prompt = f"Hub context:\n{context_text}"

    try:
        # Ask the model to use web search and return a strict JSON array of candidates.
        response = responses_client.create(
            model=settings.suggested_sources_model,
            input=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            tools=[{"type": "web_search_preview"}],
            temperature=0.2,
        )
        raw_text = _extract_response_text(response)
        # Parse defensively because the model can still wrap JSON in markdown or extra text.
        candidates = _parse_source_suggestion_candidates(raw_text)
        search_results = []
        for item in _extract_web_search_results(response)[:10]:
            search_results.append(
                {
                    "title": _get_attr(item, "title", "") or "",
                    "url": _get_attr(item, "url", "") or _get_attr(item, "link", "") or "",
                }
            )
        metadata = {
            "model": settings.suggested_sources_model,
            "usage": _extract_usage(response),
            "search_results": search_results,
        }
        return candidates, metadata
    except Exception as exc:
        logger.warning("Source suggestion discovery failed", exc_info=True)
        return [], {"error": str(exc)[:500], "model": settings.suggested_sources_model}


# Normalizes raw candidate objects into validated suggestion payloads.
def _normalize_source_suggestion_candidates(
    candidates: list[dict],
    *,
    hub_id: str,
    seed_source_ids: list[str],
    search_metadata: dict,
) -> list[dict]:
    normalized: list[dict] = []
    for candidate in candidates:
        row = _normalize_source_suggestion_candidate(
            candidate,
            hub_id=hub_id,
            seed_source_ids=seed_source_ids,
            search_metadata=search_metadata,
        )
        if row is not None:
            normalized.append(row)
    return normalized


# Validates and reshapes a single discovered suggestion candidate.
def _normalize_source_suggestion_candidate(
    candidate: dict,
    *,
    hub_id: str,
    seed_source_ids: list[str],
    search_metadata: dict,
) -> Optional[dict]:
    raw_url = str(candidate.get("url") or "").strip()
    if not raw_url:
        return None

    try:
        safe_url = _validate_public_url(raw_url)
    except Exception:
        return None

    suggested_type = str(candidate.get("type") or "web").strip().lower()
    video_id = _extract_youtube_video_id(safe_url)
    if suggested_type == "youtube" or video_id:
        if not video_id:
            return None
        return {
            "hub_id": hub_id,
            "type": "youtube",
            "status": "pending",
            "url": _canonicalize_youtube_url(video_id),
            "canonical_url": None,
            "video_id": video_id,
            "title": _trim_source_suggestion_text(candidate.get("title"), 255),
            "description": _trim_source_suggestion_text(candidate.get("description"), 1000),
            "rationale": _trim_source_suggestion_text(candidate.get("rationale"), 1000),
            "confidence": _coerce_confidence(candidate.get("confidence")),
            "seed_source_ids": seed_source_ids,
            "search_metadata": search_metadata,
        }

    canonical_url = _canonicalize_web_url(safe_url)
    if suggested_type != "web" or not canonical_url:
        return None
    return {
        "hub_id": hub_id,
        "type": "web",
        "status": "pending",
        "url": safe_url,
        "canonical_url": canonical_url,
        "video_id": None,
        "title": _trim_source_suggestion_text(candidate.get("title"), 255),
        "description": _trim_source_suggestion_text(candidate.get("description"), 1000),
        "rationale": _trim_source_suggestion_text(candidate.get("rationale"), 1000),
        "confidence": _coerce_confidence(candidate.get("confidence")),
        "seed_source_ids": seed_source_ids,
        "search_metadata": search_metadata,
    }


# Loads existing suggestion targets so duplicates can be skipped.
def _load_existing_source_suggestion_targets(client: Client, hub_id: str) -> set[tuple[str, str]]:
    rows = client.table("source_suggestions").select("type,canonical_url,video_id").eq("hub_id", hub_id).execute().data or []
    targets: set[tuple[str, str]] = set()
    for row in rows:
        key = _source_suggestion_target_key(row)
        if key is not None:
            targets.add(key)
    return targets


# Loads existing source targets so already-added content is not suggested again.
def _load_existing_source_targets(client: Client, hub_id: str) -> set[tuple[str, str]]:
    rows = client.table("sources").select("type,ingestion_metadata").eq("hub_id", hub_id).execute().data or []
    targets: set[tuple[str, str]] = set()
    for row in rows:
        source_type = str(row.get("type") or "")
        metadata = row.get("ingestion_metadata") if isinstance(row.get("ingestion_metadata"), dict) else {}
        if source_type == "youtube":
            video_id = str(metadata.get("video_id") or "")
            if not video_id:
                video_id = _extract_youtube_video_id(str(metadata.get("url") or ""))
            if video_id:
                targets.add(("youtube", video_id))
        elif source_type == "web":
            canonical_url = _canonicalize_web_url(str(metadata.get("final_url") or metadata.get("url") or ""))
            if canonical_url:
                targets.add(("web", canonical_url))
    return targets


# Deduplicates and caps candidate suggestions before insertion.
def _filter_new_source_suggestions(
    candidates: list[dict],
    *,
    existing_source_targets: set[tuple[str, str]],
    existing_suggestion_targets: set[tuple[str, str]],
    limit: int,
) -> list[dict]:
    deduped: list[dict] = []
    seen_targets = set(existing_source_targets) | set(existing_suggestion_targets)
    for candidate in candidates:
        key = _source_suggestion_target_key(candidate)
        if key is None or key in seen_targets:
            continue
        seen_targets.add(key)
        deduped.append(candidate)

    max_items = max(1, limit)
    accepted = deduped[:max_items]
    if not accepted:
        return []

    if any(str(candidate.get("type") or "").strip().lower() == "youtube" for candidate in accepted):
        return accepted

    youtube_candidate = next(
        (candidate for candidate in deduped[max_items:] if str(candidate.get("type") or "").strip().lower() == "youtube"),
        None,
    )
    if youtube_candidate is None:
        return accepted

    replacement_index = next(
        (index for index in range(len(accepted) - 1, -1, -1) if str(accepted[index].get("type") or "").strip().lower() != "youtube"),
        None,
    )
    if replacement_index is None:
        return accepted

    accepted[replacement_index] = youtube_candidate
    return accepted


# Builds a comparable key for a candidate's canonical target.
def _source_suggestion_target_key(candidate: dict) -> Optional[tuple[str, str]]:
    suggestion_type = str(candidate.get("type") or "").strip().lower()
    if suggestion_type == "youtube":
        video_id = str(candidate.get("video_id") or "").strip()
        if video_id:
            return ("youtube", video_id)
        return None
    if suggestion_type == "web":
        canonical_url = str(candidate.get("canonical_url") or "").strip()
        if canonical_url:
            return ("web", canonical_url)
    return None


# Updates the hub scan timestamp after a suggestion run finishes.
def _mark_source_suggestion_scan(client: Client, hub_id: str, *, now: datetime, generated: bool) -> None:
    payload = {"last_source_suggestion_scan_at": now.isoformat()}
    if generated:
        payload["last_source_suggestion_generated_at"] = now.isoformat()
    client.table("hubs").update(payload).eq("id", hub_id).execute()


# Extracts a YouTube video identifier from a supported video URL.
def _extract_youtube_video_id(url: str) -> Optional[str]:
    parsed = urlparse(url)
    host = (parsed.netloc or "").lower()
    if host.startswith("www."):
        host = host[4:]
    if host == "youtu.be":
        return _normalize_youtube_id(parsed.path.strip("/").split("/", 1)[0])
    if host.endswith("youtube.com") or host.endswith("youtube-nocookie.com"):
        query = parse_qs(parsed.query)
        if "v" in query and query["v"]:
            return _normalize_youtube_id(query["v"][0])
        parts = [part for part in parsed.path.split("/") if part]
        if len(parts) >= 2 and parts[0] in {"shorts", "embed", "live", "v"}:
            return _normalize_youtube_id(parts[1])
    return None


# Validates and normalizes a YouTube video identifier.
def _normalize_youtube_id(value: str) -> Optional[str]:
    cleaned = (value or "").strip()
    if not _SUGGESTION_YOUTUBE_ID_RE.fullmatch(cleaned):
        return None
    return cleaned


# Builds a canonical YouTube watch URL from a video identifier.
def _canonicalize_youtube_url(video_id: str) -> str:
    return f"https://www.youtube.com/watch?v={video_id}"


# Normalizes a web URL by removing noise such as fragments and tracking parameters.
def _canonicalize_web_url(url: str) -> Optional[str]:
    cleaned = (url or "").strip()
    if not cleaned:
        return None
    parsed = urlparse(cleaned)
    if parsed.scheme not in {"http", "https"}:
        return None
    host = (parsed.hostname or "").lower()
    if not host:
        return None
    if host.startswith("www."):
        host = host[4:]
    port = parsed.port
    if (parsed.scheme == "http" and port == 80) or (parsed.scheme == "https" and port == 443):
        port = None
    netloc = host if port is None else f"{host}:{port}"
    path = parsed.path or "/"
    path = re.sub(r"/{2,}", "/", path)
    if path != "/" and path.endswith("/"):
        path = path.rstrip("/")
    query = parse_qs(parsed.query, keep_blank_values=False)
    filtered_items: list[tuple[str, str]] = []
    for key, values in sorted(query.items()):
        normalized_key = key.lower()
        if normalized_key.startswith("utm_") or normalized_key in {"fbclid", "gclid"}:
            continue
        for value in values:
            if value:
                filtered_items.append((key, value))
    normalized_query = "&".join(f"{key}={value}" for key, value in filtered_items)
    return urlunparse((parsed.scheme.lower(), netloc, path, "", normalized_query, ""))


# Trims optional suggestion text fields to a safe maximum length.
def _trim_source_suggestion_text(value: object, max_chars: int) -> Optional[str]:
    cleaned = _normalize_text(str(value or ""))
    if not cleaned:
        return None
    if len(cleaned) <= max_chars:
        return cleaned
    return f"{cleaned[:max_chars].rstrip()}..."


# Converts a suggestion confidence value into a bounded float.
def _coerce_confidence(value: object) -> float:
    try:
        confidence = float(value)
    except (TypeError, ValueError):
        return 0.5
    return max(0.0, min(1.0, confidence))


# Parses the LLM output into a list of candidate suggestion objects.
def _parse_source_suggestion_candidates(raw: str) -> list[dict]:
    text = (raw or "").strip()
    if not text:
        return []
    if text.startswith("```"):
        lines = [line for line in text.splitlines() if not line.startswith("```")]
        text = "\n".join(lines).strip()

    candidates_to_try = [text]
    if "[" in text and "]" in text:
        candidates_to_try.append(text[text.find("[") : text.rfind("]") + 1])

    for candidate_text in candidates_to_try:
        try:
            parsed = json.loads(candidate_text)
        except json.JSONDecodeError:
            continue
        if isinstance(parsed, dict):
            parsed = parsed.get("candidates")
        if isinstance(parsed, list):
            return [item for item in parsed if isinstance(item, dict)]
    return []


# Reads an attribute from SDK response objects while keeping the caller defensive.
def _get_attr(obj: object, name: str, default: object = None) -> object:
    if isinstance(obj, dict):
        return obj.get(name, default)
    return getattr(obj, name, default)


# Extracts plain text content from an OpenAI response payload.
def _extract_response_text(response: object) -> str:
    text = _get_attr(response, "output_text")
    if isinstance(text, str) and text.strip():
        return text
    output = _get_attr(response, "output", []) or []
    for item in output:
        if _get_attr(item, "type") != "message":
            continue
        content = _get_attr(item, "content", [])
        if isinstance(content, list):
            for part in content:
                if _get_attr(part, "type") in {"output_text", "text"}:
                    part_text = _get_attr(part, "text")
                    if isinstance(part_text, str) and part_text.strip():
                        return part_text
        item_text = _get_attr(item, "text")
        if isinstance(item_text, str) and item_text.strip():
            return item_text
    return ""


# Extracts token usage information from an OpenAI response payload.
def _extract_usage(response: object) -> Optional[dict]:
    usage = _get_attr(response, "usage")
    if usage is None:
        return None
    if isinstance(usage, dict):
        return usage
    model_dump = getattr(usage, "model_dump", None)
    if callable(model_dump):
        return model_dump()
    return None


# Extracts any attached web-search result objects from an OpenAI response.
def _extract_web_search_results(response: object) -> list[object]:
    output = _get_attr(response, "output", []) or []
    results: list[object] = []
    for item in output:
        if _get_attr(item, "type") != "web_search_call":
            continue
        call = _get_attr(item, "web_search_call", item)
        call_results = _get_attr(call, "results", None)
        if call_results:
            results.extend(call_results)
    return results


# Reminder dispatch tasks and notification helpers.
# Dispatches due reminders and records any notification updates that result.
@celery_app.task(name="dispatch_reminders")
def dispatch_reminders() -> dict:
    client = _get_supabase_client()
    now = datetime.now(timezone.utc)
    lead_hours = max(1, settings.reminder_lead_hours)
    window = max(1, settings.reminder_dispatch_window_minutes)
    lead_start = now + timedelta(hours=lead_hours) - timedelta(minutes=window)
    lead_end = now + timedelta(hours=lead_hours) + timedelta(minutes=window)

    # Fetch both upcoming lead reminders and already-due reminders in the same run.
    lead_candidates = (
        client.table("reminders")
        .select("*")
        .eq("status", "scheduled")
        .gte("due_at", lead_start.isoformat())
        .lte("due_at", lead_end.isoformat())
        .execute()
        .data
    )
    due_candidates = (
        client.table("reminders")
        .select("*")
        .eq("status", "scheduled")
        .lte("due_at", now.isoformat())
        .execute()
        .data
    )

    hub_policy_cache: dict[str, dict] = {}
    sent = 0

    for reminder in lead_candidates:
        sent += _dispatch_for_reminder(
            client,
            reminder,
            "lead",
            now,
            hub_policy_cache,
        )

    for reminder in due_candidates:
        sent += _dispatch_for_reminder(
            client,
            reminder,
            "due",
            now,
            hub_policy_cache,
        )
        # Mark the reminder as sent only after the due notification path has been processed.
        _mark_reminder_sent(client, reminder["id"], now)

    return {"notifications_sent": sent}


# Processes one reminder record and sends notifications through the enabled channels.
def _dispatch_for_reminder(
    client: Client,
    reminder: dict,
    kind: str,
    now: datetime,
    hub_policy_cache: dict[str, dict],
) -> int:
    hub_id = reminder.get("hub_id")
    if not hub_id:
        return 0
    policy = _get_hub_policy(client, hub_id, hub_policy_cache)
    channels = _normalize_channels(policy.get("channels"))
    if not channels:
        return 0
    due_at = _parse_iso(reminder.get("due_at"))
    if not due_at:
        return 0
    scheduled_for = due_at
    if kind == "lead":
        if "notify_before" in reminder and reminder["notify_before"] is None:
            return 0  # User explicitly chose no notification
        notify_before = reminder.get("notify_before")
        if notify_before is not None:
            scheduled_for = due_at - timedelta(minutes=notify_before)
        else:
            lead_hours = int(policy.get("lead_hours") or settings.reminder_lead_hours)
            scheduled_for = due_at - timedelta(hours=lead_hours)
    sent = 0
    for channel in channels:
        if _create_notification_if_needed(
            client, reminder, channel, kind, scheduled_for, now
        ):
            sent += 1
    return sent


# Creates a reminder notification unless one already exists for the same dispatch window.
def _create_notification_if_needed(
    client: Client,
    reminder: dict,
    channel: str,
    kind: str,
    scheduled_for: datetime,
    now: datetime,
) -> bool:
    key = f"{reminder['id']}:{kind}:{scheduled_for.isoformat()}:{channel}"
    existing = (
        client.table("notifications")
        .select("id")
        .eq("idempotency_key", key)
        .limit(1)
        .execute()
        .data
    )
    if existing:
        return False

    payload = {
        "user_id": reminder["user_id"],
        "reminder_id": reminder["id"],
        "channel": channel,
        "status": "queued",
        "scheduled_for": scheduled_for.isoformat(),
        "idempotency_key": key,
    }
    response = client.table("notifications").insert(payload).execute()
    if not response.data:
        return False
    notification_id = response.data[0]["id"]

    # Mark as sent immediately for in-app notifications.
    _update_notification(client, notification_id, "sent", now, None, None)
    return True



# Updates an existing reminder notification with the latest dispatch outcome.
def _update_notification(
    client: Client,
    notification_id: str,
    status_value: str,
    now: datetime,
    provider_id: Optional[str],
    error: Optional[str],
) -> None:
    payload = {"status": status_value, "sent_at": now.isoformat(), "provider_id": provider_id, "error": error}
    client.table("notifications").update(payload).eq("id", notification_id).execute()


# Marks a reminder as sent after its dispatch completes successfully.
def _mark_reminder_sent(client: Client, reminder_id: str, now: datetime) -> None:
    client.table("reminders").update({"status": "sent", "sent_at": now.isoformat()}).eq("id", reminder_id).execute()


# Parses an ISO timestamp string into a timezone-aware datetime when possible.
def _parse_iso(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    cleaned = value.replace("Z", "+00:00")
    try:
        return datetime.fromisoformat(cleaned)
    except ValueError:
        return None


# Loads and caches reminder policy settings for a hub.
def _get_hub_policy(client: Client, hub_id: str, cache: dict[str, dict]) -> dict:
    cached = cache.get(hub_id)
    if cached is not None:
        return cached
    response = client.table("hubs").select("reminder_policy").eq("id", hub_id).limit(1).execute()
    policy = response.data[0].get("reminder_policy") if response.data else {}
    if not isinstance(policy, dict):
        policy = {}
    if "lead_hours" not in policy:
        policy["lead_hours"] = settings.reminder_lead_hours
    cache[hub_id] = policy
    return policy


# Normalizes a channel list into lowercase unique values.
def _normalize_channels(value: Optional[list]) -> list[str]:
    if not value:
        return ["in_app"]
    channels: list[str] = []
    for item in value:
        if not isinstance(item, str):
            continue
        key = item.lower()
        if key == "in_app":
            channels.append(key)
    return channels or ["in_app"]


# Compatibility overrides for the split worker modules.
# These definitions preserve the old `worker.tasks` surface for Celery startup and tests.
# They intentionally come last so they override the legacy helper bodies above
# until the dead implementations are removed in a later cleanup pass.
def _get_supabase_client() -> Client:
    return _common._get_supabase_client()


def _download_from_storage(storage_path: str) -> bytes:
    return _storage._download_from_storage(storage_path)


def _validate_public_url(url: str) -> str:
    return _web._validate_public_url(url)


def _ensure_public_host(hostname: str) -> None:
    return _web._ensure_public_host(hostname)


def _allowed_by_robots(url: str, user_agent: str) -> bool:
    return _web._allowed_by_robots(url, user_agent)


def _fetch_url_content(url: str) -> tuple[bytes, str, str]:
    return _web._fetch_url_content(url)


def _extract_web_text(raw: bytes, content_type: str) -> tuple[str, Optional[str]]:
    return _web._extract_web_text(raw, content_type)


def _html_to_text(html: str) -> str:
    return _web._html_to_text(html)


def _build_pseudo_doc(title: Optional[str], url: str, crawl_at: str, content_type: str, text: str) -> str:
    return _web._build_pseudo_doc(title, url, crawl_at, content_type, text)


def _fetch_youtube_transcript(
    url: str,
    language: Optional[str],
    allow_auto_captions: Optional[bool],
) -> tuple[str, dict, dict]:
    return _youtube._fetch_youtube_transcript(url, language, allow_auto_captions)


def _select_caption_track(
    info: dict,
    preferred_language: Optional[str],
    allow_auto: bool,
) -> tuple[str, str, str, str]:
    return _youtube._select_caption_track(info, preferred_language, allow_auto)


def _pick_caption_preferred(captions: dict, preferred_language: Optional[str]) -> Optional[tuple[str, str, str]]:
    return _youtube._pick_caption_preferred(captions, preferred_language)


def _pick_caption_any(captions: dict) -> Optional[tuple[str, str, str]]:
    return _youtube._pick_caption_any(captions)


def _select_caption_format(lang: str, formats: list[dict]) -> Optional[tuple[str, str, str]]:
    return _youtube._select_caption_format(lang, formats)


def _download_caption_text(url: str) -> bytes:
    return _youtube._download_caption_text(url)


def _parse_caption_text(raw: bytes, ext: str) -> str:
    return _youtube._parse_caption_text(raw, ext)


def _strip_vtt_srt(text: str) -> str:
    return _youtube._strip_vtt_srt(text)


def _strip_xml(text: str) -> str:
    return _youtube._strip_xml(text)


def _parse_json3(text: str) -> str:
    return _youtube._parse_json3(text)


def _normalize_language(value: Optional[str]) -> str:
    return _youtube._normalize_language(value)


def _format_upload_date(value: Optional[str]) -> Optional[str]:
    return _youtube._format_upload_date(value)


def _format_duration(seconds: Optional[int]) -> Optional[str]:
    return _youtube._format_duration(seconds)


def _build_youtube_pseudo_doc(info: dict, url: str, fetched_at: str, captions_meta: dict, text: str) -> str:
    return _youtube._build_youtube_pseudo_doc(info, url, fetched_at, captions_meta, text)


def _upload_pseudo_doc(client: Client, storage_path: str, content: str) -> None:
    _storage._upload_pseudo_doc(client, storage_path, content)


def _extract_text(raw: bytes, storage_path: str) -> str:
    ext = storage_path.lower()
    if ext.endswith(".pdf"):
        return _extract_pdf(raw)
    if ext.endswith(".docx"):
        return _extract_docx(raw)
    return _content._extract_text(raw, storage_path)


def _extract_pdf(raw: bytes) -> str:
    return _content._extract_pdf(raw)


def _extract_docx(raw: bytes) -> str:
    return _content._extract_docx(raw)


def _normalize_text(text: str) -> str:
    return _common._normalize_text(text)


def _trim_text(text: str, max_chars: int) -> str:
    return _common._trim_text(text, max_chars)


def _clear_existing_chunks_before(client: Client, source_id: str, cutoff: str) -> None:
    _storage._clear_existing_chunks_before(client, source_id, cutoff)


def _source_exists(client: Client, source_id: str) -> bool:
    return _storage._source_exists(client, source_id)


def _get_source_metadata(client: Client, source_id: str) -> dict:
    return _storage._get_source_metadata(client, source_id)


def _update_source(
    client: Client,
    source_id: str,
    status: str,
    failure_reason: Optional[str] = None,
    ingestion_metadata: Optional[dict] = None,
    clear_failure_reason: bool = False,
) -> None:
    _storage._update_source(
        client,
        source_id,
        status,
        failure_reason=failure_reason,
        ingestion_metadata=ingestion_metadata,
        clear_failure_reason=clear_failure_reason,
    )


def _batch(items: List, size: int) -> Iterable[List]:
    return _common._batch(items, size)


def _extract_youtube_video_id(url: str) -> Optional[str]:
    return _youtube._extract_youtube_video_id(url)


def _normalize_youtube_id(value: str) -> Optional[str]:
    return _youtube._normalize_youtube_id(value)


def _canonicalize_youtube_url(video_id: str) -> str:
    return _youtube._canonicalize_youtube_url(video_id)


def _canonicalize_web_url(url: str) -> Optional[str]:
    return _web._canonicalize_web_url(url)


def _get_attr(obj: object, name: str, default: object = None) -> object:
    return _response_utils._get_attr(obj, name, default)


def _extract_response_text(response: object) -> str:
    return _response_utils._extract_response_text(response)


def _extract_usage(response: object) -> Optional[dict]:
    return _response_utils._extract_usage(response)


def _extract_web_search_results(response: object) -> list[object]:
    return _response_utils._extract_web_search_results(response)


def _parse_iso(value: Optional[str]) -> Optional[datetime]:
    return _common._parse_iso(value)
